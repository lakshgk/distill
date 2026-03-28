"""
distill.parsers.docx
~~~~~~~~~~~~~~~~~~~~
Parser for Microsoft Word documents (.docx, .doc).

Primary path:   mammoth  (.docx → clean HTML → IR)
Metadata path:  python-docx (document core + app properties)
Legacy path:    LibreOffice headless (.doc → .docx pre-conversion)
Fallback:       pandoc  (complex documents, tracked changes)

Install:
    pip install distill-core           # includes mammoth, python-docx
    pip install distill-core[legacy]   # adds LibreOffice support for .doc
"""

from __future__ import annotations

import io
import re
import subprocess
import zipfile
from pathlib import Path
from typing import Optional, Union

import defusedxml.ElementTree as ET

from distill.ir import (
    BlockQuote, CodeBlock, Document, DocumentMetadata, Image, ImageType,
    List, ListItem, Paragraph, Section, Table, TableCell, TableRow, TextRun,
)
from distill.parsers.base import ParseError, ParseOptions, Parser, UnsupportedFormatError
from distill.registry import registry


# ── HTML → IR helpers ────────────────────────────────────────────────────────

def _parse_runs(
    elem,
    bold: bool = False,
    italic: bool = False,
    code: bool = False,
    strikethrough: bool = False,
    href: Optional[str] = None,
    skip_tags: Optional[set] = None,
) -> list[TextRun]:
    """
    Recursively extract TextRun objects from an HTML element,
    propagating inline formatting down through nested tags.
    """
    skip_tags = skip_tags or set()
    runs: list[TextRun] = []

    if elem.text:
        runs.append(TextRun(
            text=elem.text,
            bold=bold, italic=italic, code=code,
            strikethrough=strikethrough, href=href,
        ))

    for child in elem:
        tag = child.tag.lower()

        if tag in skip_tags:
            if child.tail:
                runs.append(TextRun(
                    text=child.tail,
                    bold=bold, italic=italic, code=code,
                    strikethrough=strikethrough, href=href,
                ))
            continue

        c_bold        = bold or tag in ("strong", "b")
        c_italic      = italic or tag in ("em", "i")
        c_code        = code or tag == "code"
        c_strike      = strikethrough or tag in ("del", "s", "strike")
        c_href        = child.get("href") if tag == "a" else href

        runs.extend(_parse_runs(
            child,
            bold=c_bold, italic=c_italic, code=c_code,
            strikethrough=c_strike, href=c_href,
            skip_tags=skip_tags,
        ))

        if child.tail:
            runs.append(TextRun(
                text=child.tail,
                bold=bold, italic=italic, code=code,
                strikethrough=strikethrough, href=href,
            ))

    return [r for r in runs if r.text]


def _parse_list(elem) -> List:
    """Parse a <ul> or <ol> element into an IR List."""
    ordered = elem.tag.lower() == "ol"
    items: list[ListItem] = []

    for li in elem:
        if li.tag.lower() != "li":
            continue

        nested: list[List] = []
        for child in li:
            if child.tag.lower() in ("ul", "ol"):
                nested.append(_parse_list(child))

        content_runs = _parse_runs(li, skip_tags={"ul", "ol"})
        items.append(ListItem(content=content_runs, children=nested))

    return List(items=items, ordered=ordered)


def _parse_table(elem) -> Table:
    """Parse a <table> element into an IR Table."""
    rows: list[TableRow] = []

    for child in elem:
        tag = child.tag.lower()
        if tag in ("thead", "tbody", "tfoot"):
            is_header_section = tag == "thead"
            for tr in child:
                if tr.tag.lower() == "tr":
                    rows.append(_parse_row(tr, header_section=is_header_section))
        elif tag == "tr":
            rows.append(_parse_row(child))

    return Table(rows=rows)


def _parse_row(elem, header_section: bool = False) -> TableRow:
    cells: list[TableCell] = []
    for cell in elem:
        tag = cell.tag.lower()
        if tag not in ("td", "th"):
            continue
        is_header = tag == "th" or header_section
        runs = _parse_runs(cell)
        colspan = int(cell.get("colspan", 1) or 1)
        rowspan = int(cell.get("rowspan", 1) or 1)
        cells.append(TableCell(
            content=[Paragraph(runs=runs)] if runs else [],
            is_header=is_header,
            colspan=colspan,
            rowspan=rowspan,
        ))
    return TableRow(cells=cells)


def _parse_block(elem):
    """Parse a single block-level HTML element into an IR block node."""
    tag = elem.tag.lower()

    if tag == "p":
        runs = _parse_runs(elem)
        return Paragraph(runs=runs) if runs else None

    if tag in ("ul", "ol"):
        return _parse_list(elem)

    if tag == "table":
        return _parse_table(elem)

    if tag == "pre":
        code_text = "".join(elem.itertext())
        lang = None
        code_child = elem.find("code")
        if code_child is not None:
            cls = code_child.get("class", "")
            if cls.startswith("language-"):
                lang = cls[9:]
        return CodeBlock(code=code_text.strip(), language=lang)

    if tag == "blockquote":
        blocks = [_parse_block(c) for c in elem]
        blocks = [b for b in blocks if b is not None]
        return BlockQuote(content=blocks) if blocks else None

    if tag == "img":
        alt = elem.get("alt") or None
        return Image(image_type=ImageType.UNKNOWN, alt_text=alt)

    return None


def _html_to_sections(html: str) -> list[Section]:
    """
    Convert mammoth's HTML output into a list of IR Sections.

    Heading tags (h1–h6) create new sections; all other block elements
    are appended as blocks to the current section.
    """
    try:
        root = ET.fromstring(f"<root>{html}</root>")
    except ET.ParseError:
        # If XML parsing fails, sanitise and retry
        clean = re.sub(r"&(?!(amp|lt|gt|quot|apos|#\d+|#x[0-9a-fA-F]+);)", "&amp;", html)
        root = ET.fromstring(f"<root>{clean}</root>")

    sections: list[Section] = []
    current: Optional[Section] = None

    for elem in root:
        tag = elem.tag.lower()

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            heading_runs = _parse_runs(elem)
            current = Section(heading=heading_runs or None, level=level)
            sections.append(current)
        else:
            if current is None:
                current = Section(level=0)
                sections.append(current)
            block = _parse_block(elem)
            if block is not None:
                current.blocks.append(block)

    return sections


# ── Metadata extraction ───────────────────────────────────────────────────────

_APP_PROPS_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"


def _extract_metadata(path: Path) -> DocumentMetadata:
    """
    Extract all available metadata from a .docx file via python-docx.
    Falls back gracefully if any field is unavailable.
    """
    import docx as python_docx

    meta = DocumentMetadata(source_format="docx", source_path=str(path))

    try:
        doc_obj    = python_docx.Document(str(path))
        core_props = doc_obj.core_properties

        meta.title       = core_props.title or None
        meta.author      = core_props.author or None
        meta.subject     = core_props.subject or None
        meta.description = core_props.comments or None   # dc:description in OOXML
        meta.created_at  = core_props.created.isoformat() if core_props.created else None
        meta.modified_at = core_props.modified.isoformat() if core_props.modified else None

        if core_props.keywords:
            meta.keywords = [
                k.strip()
                for k in re.split(r"[,;]", core_props.keywords)
                if k.strip()
            ]

        meta.word_count = sum(len(p.text.split()) for p in doc_obj.paragraphs)

        # Page count lives in app properties (not core properties)
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT
            app_part = doc_obj.part.package.part_related_by(RT.EXTENDED_PROPERTIES)
            app_tree = ET.fromstring(app_part.blob)
            pages_elem = app_tree.find(f"{{{_APP_PROPS_NS}}}Pages")
            if pages_elem is not None and pages_elem.text:
                meta.page_count = int(pages_elem.text)
        except Exception:
            pass

    except Exception:
        pass  # metadata is always best-effort

    return meta


# ── Pandoc fallback ───────────────────────────────────────────────────────────

def _pandoc_fallback(path: Path) -> Optional[str]:
    """
    Call pandoc to convert .docx → GFM Markdown.
    Returns the Markdown string, or None if pandoc is unavailable or fails.
    """
    try:
        result = subprocess.run(
            ["pandoc", "--from=docx", "--to=gfm", str(path)],
            capture_output=True, text=True, timeout=30,
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return None


# ── Security constants ────────────────────────────────────────────────────────

_MAX_FILE_BYTES      = 50  * 1024 * 1024   # 50 MB compressed input limit
_MAX_UNZIP_BYTES     = 500 * 1024 * 1024   # 500 MB uncompressed zip-bomb limit


def _check_input_size(source: Union[str, "Path", bytes], max_bytes: int) -> None:
    """Raise ParseError if the input exceeds the configured size limit."""
    if isinstance(source, bytes):
        size = len(source)
    else:
        try:
            size = Path(source).stat().st_size
        except OSError:
            return  # can't stat — let the parser surface any real error
    if size > max_bytes:
        mb = max_bytes // (1024 * 1024)
        raise ParseError(
            f"Input file exceeds the {mb} MB size limit "
            f"({size / 1024 / 1024:.1f} MB). "
            f"Increase the limit via options.extra['max_file_size']."
        )


def _check_zip_bomb(source: Union[str, "Path", bytes], max_unzip_bytes: int) -> None:
    """
    Inspect a DOCX (ZIP) archive and raise ParseError if the total
    uncompressed size exceeds max_unzip_bytes (default 500 MB).
    This catches zip-bomb payloads before any extraction occurs.
    """
    try:
        data = source if isinstance(source, bytes) else open(source, "rb").read()
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            total = sum(info.file_size for info in zf.infolist())
        if total > max_unzip_bytes:
            mb = max_unzip_bytes // (1024 * 1024)
            raise ParseError(
                f"DOCX archive uncompressed size ({total / 1024 / 1024:.0f} MB) "
                f"exceeds the {mb} MB safety limit. "
                f"Increase via options.extra['max_unzip_size']."
            )
    except zipfile.BadZipFile:
        raise ParseError("File is not a valid DOCX (ZIP) archive.")
    except ParseError:
        raise
    except Exception:
        pass  # any other error will surface during actual parsing


# ── Parser ────────────────────────────────────────────────────────────────────

@registry.register
class DocxParser(Parser):
    """
    Parses .docx files using mammoth for content and python-docx for metadata.
    Falls back to pandoc for complex documents where mammoth yields empty content.
    """

    extensions = [".docx"]
    mime_types = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ]
    requires          = ["mammoth", "docx"]
    optional_requires = ["pandoc"]

    def parse(
        self,
        source: Union[str, Path, bytes],
        options: Optional[ParseOptions] = None,
    ) -> Document:
        options = options or ParseOptions()

        try:
            import mammoth
        except ImportError as e:
            raise ParseError(f"Missing dependency: {e}") from e

        # ── Security checks ───────────────────────────────────────────────────
        max_file  = options.extra.get("max_file_size",  _MAX_FILE_BYTES)
        max_unzip = options.extra.get("max_unzip_size", _MAX_UNZIP_BYTES)
        _check_input_size(source, max_file)
        _check_zip_bomb(source, max_unzip)

        path = Path(source) if not isinstance(source, bytes) else None

        # ── Metadata ─────────────────────────────────────────────────────────
        metadata = _extract_metadata(path) if path else DocumentMetadata(source_format="docx")

        # ── mammoth: .docx → HTML ─────────────────────────────────────────────
        try:
            if path:
                with open(path, "rb") as f:
                    result = mammoth.convert_to_html(f)
            else:
                import io
                result = mammoth.convert_to_html(io.BytesIO(source))
        except Exception as e:
            raise ParseError(f"mammoth failed: {e}") from e

        html     = result.value
        document = Document(metadata=metadata)

        for msg in result.messages:
            document.warnings.append(f"[docx] {msg}")

        # ── HTML → IR ────────────────────────────────────────────────────────
        sections = _html_to_sections(html)

        # Check if mammoth produced meaningful content
        has_content = any(
            s.blocks or s.heading for s in sections
        )

        if not has_content and path:
            # Pandoc fallback: wrap GFM output in a single section paragraph
            md = _pandoc_fallback(path)
            if md:
                document.sections.append(
                    Section(level=0, blocks=[Paragraph(runs=[TextRun(text=md)])])
                )
                document.warnings.append(
                    "[docx] mammoth produced no content; pandoc fallback used"
                )
                return document
            document.warnings.append(
                "[docx] mammoth produced no content and pandoc is unavailable"
            )

        document.sections.extend(sections)
        return document


@registry.register
class DocLegacyParser(Parser):
    """
    Converts legacy .doc files to .docx via LibreOffice headless, then
    delegates to DocxParser for the actual content extraction.

    Requires LibreOffice to be installed and on PATH (or DISTILL_LIBREOFFICE
    environment variable set to the full binary path).
    """

    extensions            = [".doc"]
    mime_types            = ["application/msword"]
    requires              = ["mammoth", "docx"]
    requires_libreoffice  = True

    def parse(
        self,
        source: Union[str, Path, bytes],
        options: Optional[ParseOptions] = None,
    ) -> Document:
        import shutil
        from distill.parsers._libreoffice import convert_via_libreoffice

        options = options or ParseOptions()
        timeout = options.extra.get("libreoffice_timeout", 60)

        output_path = convert_via_libreoffice(source, "docx", timeout=timeout)
        try:
            doc = DocxParser().parse(output_path, options)
            # Preserve the original source path/format in metadata
            doc.metadata.source_format = "doc"
            if not isinstance(source, bytes):
                doc.metadata.source_path = str(source)
            return doc
        finally:
            shutil.rmtree(output_path.parent, ignore_errors=True)
