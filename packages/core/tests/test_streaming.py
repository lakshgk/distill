"""
tests/test_streaming.py
~~~~~~~~~~~~~~~~~~~~~~~
Tests for the streaming API: MarkdownRenderer.render_stream(),
Document.render_stream(), and convert_stream().
"""

from __future__ import annotations

import types

import pytest

from distill import convert_stream
from distill.ir import (
    BlockQuote,
    CodeBlock,
    Document,
    DocumentMetadata,
    List,
    ListItem,
    Paragraph,
    Section,
    Table,
    TableCell,
    TableRow,
    TextRun,
)
from distill.renderer import MarkdownRenderer


# ── Helpers ──────────────────────────────────────────────────────────────────

def _doc(*headings: str) -> Document:
    """Build a Document with one section per heading, each containing one paragraph."""
    return Document(
        metadata=DocumentMetadata(title="Test", author="Tester"),
        sections=[
            Section(
                level=1,
                heading=[TextRun(h)],
                blocks=[Paragraph([TextRun(f"Content of {h}.")])],
            )
            for h in headings
        ],
    )


def _empty_doc() -> Document:
    return Document()


# ── MarkdownRenderer.render_stream() ─────────────────────────────────────────

class TestRenderStream:

    def test_returns_generator(self):
        doc = _doc("A")
        result = MarkdownRenderer().render_stream(doc)
        assert isinstance(result, types.GeneratorType)

    def test_yields_one_chunk_per_section(self):
        doc = _doc("Alpha", "Beta", "Gamma")
        chunks = list(MarkdownRenderer().render_stream(doc))
        assert len(chunks) == 3

    def test_each_chunk_contains_heading(self):
        doc = _doc("Alpha", "Beta")
        chunks = list(MarkdownRenderer().render_stream(doc))
        assert "# Alpha" in chunks[0]
        assert "# Beta"  in chunks[1]

    def test_each_chunk_contains_paragraph(self):
        doc = _doc("Alpha")
        chunks = list(MarkdownRenderer().render_stream(doc))
        assert "Content of Alpha." in chunks[0]

    def test_empty_sections_skipped(self):
        doc = Document(sections=[
            Section(level=1, heading=[TextRun("Non-empty")], blocks=[Paragraph([TextRun("hi")])]),
            Section(level=1, heading=None, blocks=[]),  # empty — no heading, no blocks
            Section(level=1, heading=[TextRun("Also non-empty")], blocks=[Paragraph([TextRun("bye")])]),
        ])
        chunks = list(MarkdownRenderer().render_stream(doc))
        assert len(chunks) == 2

    def test_empty_document_yields_nothing(self):
        chunks = list(MarkdownRenderer().render_stream(_empty_doc()))
        assert chunks == []

    def test_front_matter_is_first_chunk(self):
        doc = _doc("Section One")
        chunks = list(MarkdownRenderer(front_matter=True).render_stream(doc))
        assert chunks[0].startswith("---")
        assert "title:" in chunks[0]
        assert "# Section One" in chunks[1]

    def test_front_matter_chunk_count(self):
        doc = _doc("A", "B")
        chunks = list(MarkdownRenderer(front_matter=True).render_stream(doc))
        # front matter + 2 sections
        assert len(chunks) == 3

    def test_no_front_matter_when_disabled(self):
        doc = _doc("A")
        chunks = list(MarkdownRenderer(front_matter=False).render_stream(doc))
        assert not chunks[0].startswith("---")

    def test_front_matter_omitted_when_metadata_empty(self):
        doc = Document(sections=[
            Section(level=1, heading=[TextRun("Hi")], blocks=[Paragraph([TextRun("x")])])
        ])
        chunks = list(MarkdownRenderer(front_matter=True).render_stream(doc))
        # No metadata fields populated → no front matter block emitted
        assert len(chunks) == 1
        assert "# Hi" in chunks[0]

    def test_joined_chunks_equal_render_output(self):
        doc = _doc("Alpha", "Beta", "Gamma")
        renderer = MarkdownRenderer()
        full   = renderer.render(doc)
        joined = "\n\n".join(renderer.render_stream(doc))
        assert full == joined

    def test_joined_chunks_equal_render_output_with_front_matter(self):
        doc = _doc("Alpha", "Beta")
        renderer = MarkdownRenderer(front_matter=True)
        full   = renderer.render(doc)
        joined = "\n\n".join(renderer.render_stream(doc))
        assert full == joined

    def test_chunks_are_strings(self):
        doc = _doc("A", "B")
        for chunk in MarkdownRenderer().render_stream(doc):
            assert isinstance(chunk, str)

    def test_table_in_section_rendered_in_chunk(self):
        table = Table(rows=[
            TableRow(cells=[TableCell(content=[TextRun("H1")]), TableCell(content=[TextRun("H2")])]),
            TableRow(cells=[TableCell(content=[TextRun("v1")]), TableCell(content=[TextRun("v2")])]),
        ])
        doc = Document(sections=[
            Section(level=1, heading=[TextRun("Data")], blocks=[table])
        ])
        chunks = list(MarkdownRenderer().render_stream(doc))
        assert len(chunks) == 1
        assert "|" in chunks[0]

    def test_code_block_in_chunk(self):
        doc = Document(sections=[
            Section(level=1, heading=[TextRun("Code")], blocks=[CodeBlock(code="x = 1", language="python")])
        ])
        chunks = list(MarkdownRenderer().render_stream(doc))
        assert "```python" in chunks[0]

    def test_blockquote_in_chunk(self):
        bq = BlockQuote(content=[Paragraph([TextRun("A quote.")])])
        doc = Document(sections=[
            Section(level=1, heading=[TextRun("Notes")], blocks=[bq])
        ])
        chunks = list(MarkdownRenderer().render_stream(doc))
        assert "> A quote." in chunks[0]

    def test_list_in_chunk(self):
        lst = List(items=[ListItem(content=[TextRun("item one")]), ListItem(content=[TextRun("item two")])])
        doc = Document(sections=[
            Section(level=1, heading=[TextRun("Items")], blocks=[lst])
        ])
        chunks = list(MarkdownRenderer().render_stream(doc))
        assert "- item one" in chunks[0]

    def test_max_heading_depth_respected(self):
        doc = Document(sections=[
            Section(level=6, heading=[TextRun("Deep")], blocks=[Paragraph([TextRun("x")])])
        ])
        renderer = MarkdownRenderer(max_heading_depth=2)
        chunks = list(renderer.render_stream(doc))
        assert chunks[0].startswith("## Deep")


# ── Document.render_stream() convenience method ───────────────────────────────

class TestDocumentRenderStream:

    def test_convenience_method_yields_same_as_renderer(self):
        doc = _doc("X", "Y")
        from_convenience = list(doc.render_stream())
        from_renderer    = list(MarkdownRenderer().render_stream(doc))
        assert from_convenience == from_renderer

    def test_convenience_method_with_front_matter(self):
        doc = _doc("X")
        chunks = list(doc.render_stream(front_matter=True))
        assert chunks[0].startswith("---")

    def test_convenience_method_is_generator(self):
        doc = _doc("X")
        result = doc.render_stream()
        assert isinstance(result, types.GeneratorType)


# ── convert_stream() ─────────────────────────────────────────────────────────

class TestConvertStream:

    @pytest.fixture()
    def docx_path(self, tmp_path):
        """Write a minimal two-section .docx to a temp file and return its path."""
        import io
        import docx as pydocx
        buf = io.BytesIO()
        d = pydocx.Document()
        d.add_heading("Chapter One", level=1)
        d.add_paragraph("First paragraph.")
        d.add_heading("Chapter Two", level=1)
        d.add_paragraph("Second paragraph.")
        d.save(buf)
        p = tmp_path / "sample.docx"
        p.write_bytes(buf.getvalue())
        return p

    def test_returns_generator(self, docx_path):
        result = convert_stream(docx_path)
        assert isinstance(result, types.GeneratorType)

    def test_yields_strings(self, docx_path):
        for chunk in convert_stream(docx_path):
            assert isinstance(chunk, str)

    def test_yields_multiple_chunks_for_multi_section_doc(self, docx_path):
        chunks = list(convert_stream(docx_path))
        assert len(chunks) >= 2

    def test_chunks_contain_headings(self, docx_path):
        chunks = list(convert_stream(docx_path))
        combined = "\n".join(chunks)
        assert "Chapter One" in combined
        assert "Chapter Two" in combined

    def test_include_metadata_emits_front_matter_first(self, docx_path):
        chunks = list(convert_stream(docx_path, include_metadata=True))
        assert chunks[0].startswith("---")

    def test_joined_matches_convert_output(self, docx_path):
        from distill import convert
        streamed = "\n\n".join(convert_stream(docx_path))
        full     = convert(docx_path).markdown
        assert streamed == full

    def test_joined_matches_convert_output_with_metadata(self, docx_path):
        from distill import convert
        streamed = "\n\n".join(convert_stream(docx_path, include_metadata=True))
        full     = convert(docx_path, include_metadata=True).markdown
        assert streamed == full

    def test_is_lazy_generator(self, docx_path):
        """convert_stream must not materialise all output before the first next() call."""
        gen   = convert_stream(docx_path)
        first = next(gen)
        assert isinstance(first, str)
