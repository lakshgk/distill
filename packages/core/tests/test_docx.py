"""
Tests for distill.parsers.docx — DocxParser.

Fixtures are built programmatically via python-docx so the test suite
has no binary blobs checked into the repository.
"""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest

from distill.ir import (
    CodeBlock, Document, List, Paragraph, Section, Table, TextRun,
)
from distill.parsers.base import ParseError
from distill.parsers.docx import (
    DocxParser,
    _check_input_size,
    _check_zip_bomb,
)


# ── Fixture helpers ──────────────────────────────────────────────────────────

def _make_docx(
    *,
    title: str = "Test Doc",
    author: str = "Test Author",
    subject: str = "",
    comments: str = "",
    keywords: str = "",
    paragraphs: list[str] | None = None,
    headings: list[tuple[str, int]] | None = None,
) -> bytes:
    """Build a minimal .docx in memory and return its raw bytes."""
    import docx as python_docx

    buf = io.BytesIO()
    doc = python_docx.Document()

    # Core properties
    cp = doc.core_properties
    cp.title    = title
    cp.author   = author
    cp.subject  = subject
    cp.comments = comments
    cp.keywords = keywords

    for text, level in (headings or []):
        doc.add_heading(text, level=level)

    for text in (paragraphs or ["Hello world."]):
        doc.add_paragraph(text)

    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_table() -> bytes:
    """Build a .docx that contains a simple 2×2 table."""
    import docx as python_docx

    buf = io.BytesIO()
    doc = python_docx.Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Name"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Alpha"
    table.cell(1, 1).text = "1"
    doc.save(buf)
    return buf.getvalue()


def _make_bad_zip() -> bytes:
    """Return bytes that look like a ZIP but are corrupt."""
    return b"PK\x03\x04" + b"\x00" * 100


# ── Parser availability ───────────────────────────────────────────────────────

class TestParserAvailability:
    def test_is_available(self):
        assert DocxParser.is_available()

    def test_extensions(self):
        assert ".docx" in DocxParser.extensions

    def test_missing_requires_empty(self):
        assert DocxParser.missing_requires() == []


# ── Basic parsing ─────────────────────────────────────────────────────────────

class TestBasicParsing:
    def test_returns_document(self):
        data = _make_docx(paragraphs=["Simple paragraph."])
        result = DocxParser().parse(data)
        assert isinstance(result, Document)

    def test_paragraph_extracted(self):
        data = _make_docx(paragraphs=["Hello from Distill."])
        doc  = DocxParser().parse(data)
        text = _all_text(doc)
        assert "Hello from Distill." in text

    def test_heading_creates_section(self):
        data = _make_docx(
            headings=[("Introduction", 1), ("Background", 2)],
            paragraphs=["Body text."],
        )
        doc = DocxParser().parse(data)
        headings = _heading_texts(doc)
        assert "Introduction" in headings
        assert "Background" in headings

    def test_multiple_paragraphs(self):
        paras = ["First paragraph.", "Second paragraph.", "Third paragraph."]
        data  = _make_docx(paragraphs=paras)
        doc   = DocxParser().parse(data)
        text  = _all_text(doc)
        for p in paras:
            assert p in text

    def test_table_extracted(self):
        data = _make_docx_with_table()
        doc  = DocxParser().parse(data)
        tables = _collect_blocks(doc, Table)
        assert len(tables) >= 1
        cell_texts = [
            run.text
            for tbl in tables
            for row in tbl.rows
            for cell in row.cells
            for para in cell.content
            if isinstance(para, Paragraph)
            for run in para.runs
        ]
        assert "Name" in cell_texts
        assert "Alpha" in cell_texts

    def test_empty_doc_no_crash(self):
        data = _make_docx(paragraphs=[], headings=[])
        doc  = DocxParser().parse(data)
        assert isinstance(doc, Document)

    def test_accepts_path(self, tmp_path):
        p = tmp_path / "sample.docx"
        p.write_bytes(_make_docx(paragraphs=["Path-based test."]))
        doc = DocxParser().parse(str(p))
        assert "Path-based test." in _all_text(doc)

    def test_accepts_path_object(self, tmp_path):
        p = tmp_path / "sample.docx"
        p.write_bytes(_make_docx(paragraphs=["Path object test."]))
        doc = DocxParser().parse(p)
        assert "Path object test." in _all_text(doc)


# ── Metadata extraction ────────────────────────────────────────────────────────
#
# Metadata is only populated when parsing from a file path (not raw bytes),
# because python-docx needs a real file to read core properties.
# All metadata tests therefore write the fixture to a tmp_path first.

class TestMetadata:
    def test_title(self, tmp_path):
        p = tmp_path / "t.docx"
        p.write_bytes(_make_docx(title="My Report"))
        doc = DocxParser().parse(p)
        assert doc.metadata.title == "My Report"

    def test_author(self, tmp_path):
        p = tmp_path / "t.docx"
        p.write_bytes(_make_docx(author="Jane Doe"))
        doc = DocxParser().parse(p)
        assert doc.metadata.author == "Jane Doe"

    def test_subject(self, tmp_path):
        p = tmp_path / "t.docx"
        p.write_bytes(_make_docx(subject="Financial Analysis"))
        doc = DocxParser().parse(p)
        assert doc.metadata.subject == "Financial Analysis"

    def test_description_from_comments(self, tmp_path):
        p = tmp_path / "t.docx"
        p.write_bytes(_make_docx(comments="Quarterly earnings review"))
        doc = DocxParser().parse(p)
        assert doc.metadata.description == "Quarterly earnings review"

    def test_keywords_comma_separated(self, tmp_path):
        p = tmp_path / "t.docx"
        p.write_bytes(_make_docx(keywords="finance, earnings, Q1"))
        doc = DocxParser().parse(p)
        assert "finance" in doc.metadata.keywords
        assert "earnings" in doc.metadata.keywords
        assert "Q1" in doc.metadata.keywords

    def test_keywords_semicolon_separated(self, tmp_path):
        p = tmp_path / "t.docx"
        p.write_bytes(_make_docx(keywords="alpha;beta;gamma"))
        doc = DocxParser().parse(p)
        assert doc.metadata.keywords == ["alpha", "beta", "gamma"]

    def test_source_format(self):
        # source_format is always set regardless of bytes vs path
        data = _make_docx()
        doc  = DocxParser().parse(data)
        assert doc.metadata.source_format == "docx"

    def test_word_count_positive(self, tmp_path):
        p = tmp_path / "t.docx"
        p.write_bytes(_make_docx(paragraphs=["One two three four five."]))
        doc = DocxParser().parse(p)
        assert doc.metadata.word_count is not None
        assert doc.metadata.word_count > 0


# ── Security checks ────────────────────────────────────────────────────────────

class TestSecurity:
    def test_input_size_check_bytes(self):
        oversized = b"x" * (55 * 1024 * 1024)
        with pytest.raises(ParseError, match="50 MB"):
            _check_input_size(oversized, 50 * 1024 * 1024)

    def test_input_size_check_path(self, tmp_path):
        big = tmp_path / "big.bin"
        big.write_bytes(b"x" * (55 * 1024 * 1024))
        with pytest.raises(ParseError, match="50 MB"):
            _check_input_size(str(big), 50 * 1024 * 1024)

    def test_input_size_ok(self):
        small = b"x" * 100
        _check_input_size(small, 50 * 1024 * 1024)  # must not raise

    def test_zip_bomb_detected(self):
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("big.txt", "A" * (501 * 1024 * 1024))
        with pytest.raises(ParseError, match="500 MB"):
            _check_zip_bomb(buf.getvalue(), 500 * 1024 * 1024)

    def test_bad_zip_raises_parse_error(self):
        with pytest.raises(ParseError, match="not a valid DOCX"):
            _check_zip_bomb(b"not a zip file at all", 500 * 1024 * 1024)

    def test_parser_rejects_oversized_file(self):
        oversized = b"x" * (55 * 1024 * 1024)
        with pytest.raises(ParseError, match="50 MB"):
            DocxParser().parse(oversized)

    def test_custom_size_limit_via_options(self):
        from distill.parsers.base import ParseOptions
        # 10 MB limit — a 15 MB blob should fail
        data = b"x" * (15 * 1024 * 1024)
        opts = ParseOptions(extra={"max_file_size": 10 * 1024 * 1024})
        with pytest.raises(ParseError, match="10 MB"):
            DocxParser().parse(data, options=opts)

    def test_defusedxml_used(self):
        import defusedxml.ElementTree
        from distill.parsers import docx as docx_mod
        assert docx_mod.ET is defusedxml.ElementTree


# ── Render integration ─────────────────────────────────────────────────────────

class TestRenderIntegration:
    def test_front_matter_in_markdown(self, tmp_path):
        p = tmp_path / "t.docx"
        p.write_bytes(_make_docx(title="My Report", author="Alice"))
        doc = DocxParser().parse(p)
        md  = doc.render(front_matter=True)
        assert "---" in md
        assert "My Report" in md
        assert "Alice" in md

    def test_no_front_matter_when_suppressed(self):
        data = _make_docx(title="My Report")
        doc  = DocxParser().parse(data)
        md   = doc.render(front_matter=False)
        assert "---" not in md

    def test_headings_render_as_markdown(self):
        data = _make_docx(headings=[("Overview", 1)], paragraphs=["Content."])
        doc  = DocxParser().parse(data)
        md   = doc.render(front_matter=False)
        assert "# Overview" in md

    def test_warnings_list_is_present(self):
        data = _make_docx()
        doc  = DocxParser().parse(data)
        assert isinstance(doc.warnings, list)


# ── Helpers ────────────────────────────────────────────────────────────────────

def _all_text(doc: Document) -> str:
    parts: list[str] = []
    for section in doc.sections:
        for block in section.blocks:
            if isinstance(block, Paragraph):
                for run in block.runs:
                    parts.append(run.text)
    return " ".join(parts)


def _heading_texts(doc: Document) -> list[str]:
    result: list[str] = []
    for section in doc.sections:
        if section.heading:
            result.extend(r.text for r in section.heading)
    return result


def _collect_blocks(doc: Document, block_type):
    result = []
    for section in doc.sections:
        for block in section.blocks:
            if isinstance(block, block_type):
                result.append(block)
    return result
