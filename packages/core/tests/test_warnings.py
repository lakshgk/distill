"""
Tests for distill.warnings — WarningCollector and ConversionWarning.
"""

from __future__ import annotations

import io
import tempfile
from pathlib import Path

import pytest

from distill.warnings import ConversionWarning, WarningCollector, WarningType


# ── WarningCollector unit tests ──────────────────────────────────────────────

def test_warning_type_table_complex_exists():
    assert WarningType.table_complex.value == "table_complex"


def test_warning_type_vision_caption_failed_exists():
    assert WarningType.vision_caption_failed.value == "vision_caption_failed"


def test_add_and_all_accumulate():
    c = WarningCollector()
    c.add(ConversionWarning(type=WarningType.MATH_DETECTED, message="found math"))
    c.add(ConversionWarning(type=WarningType.TABLE_TRUNCATED, message="rows capped"))
    assert len(c.all()) == 2


def test_all_returns_copy():
    c = WarningCollector()
    c.add(ConversionWarning(type=WarningType.SCANNED_CONTENT, message="scanned"))
    snapshot = c.all()
    c.add(ConversionWarning(type=WarningType.TABLE_TRUNCATED, message="truncated"))
    assert len(snapshot) == 1  # original snapshot unaffected


def test_has_returns_true_when_present():
    c = WarningCollector()
    c.add(ConversionWarning(type=WarningType.MATH_DETECTED, message="math"))
    assert c.has(WarningType.MATH_DETECTED) is True


def test_has_returns_false_when_absent():
    c = WarningCollector()
    c.add(ConversionWarning(type=WarningType.MATH_DETECTED, message="math"))
    assert c.has(WarningType.SCANNED_CONTENT) is False


def test_has_on_empty_collector():
    c = WarningCollector()
    assert c.has(WarningType.TABLE_TRUNCATED) is False


# ── to_dict serialisation ────────────────────────────────────────────────────

def test_to_dict_omits_none_optional_fields():
    c = WarningCollector()
    c.add(ConversionWarning(type=WarningType.CROSS_PAGE_TABLE, message="spans page"))
    d = c.to_dict()
    assert len(d) == 1
    assert "pages" not in d[0]
    assert "count" not in d[0]


def test_to_dict_includes_pages_and_count_when_set():
    c = WarningCollector()
    c.add(ConversionWarning(
        type=WarningType.MATH_DETECTED,
        message="14 equations",
        count=14,
        pages=[2, 3, 7],
    ))
    d = c.to_dict()
    assert d[0]["count"] == 14
    assert d[0]["pages"] == [2, 3, 7]


def test_to_dict_type_is_string_value():
    c = WarningCollector()
    c.add(ConversionWarning(type=WarningType.AUDIO_QUALITY_LOW, message="low bitrate"))
    d = c.to_dict()
    assert d[0]["type"] == "audio_quality_low"


def test_to_dict_on_empty_collector_returns_empty_list():
    c = WarningCollector()
    assert c.to_dict() == []


def test_to_dict_preserves_order():
    c = WarningCollector()
    c.add(ConversionWarning(type=WarningType.MATH_DETECTED, message="first"))
    c.add(ConversionWarning(type=WarningType.TABLE_TRUNCATED, message="second"))
    d = c.to_dict()
    assert d[0]["type"] == "math_detected"
    assert d[1]["type"] == "table_truncated"


# ── Integration: convert() returns structured_warnings ──────────────────────

def _make_docx_bytes() -> bytes:
    import docx as _docx
    doc = _docx.Document()
    doc.add_heading("Test", level=1)
    doc.add_paragraph("Body text.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_warning_type_image_write_failed_exists():
    assert WarningType.image_write_failed.value == "image_write_failed"


def test_extract_image_unwritable_emits_warning():
    from distill.parsers.base import extract_image

    collector = WarningCollector()
    result = extract_image(
        image_bytes=b"fakepng",
        ext="png",
        image_dir=Path("Z:/nonexistent_drive/images"),
        filename="test_0",
        collector=collector,
    )
    assert result is None
    assert collector.has(WarningType.image_write_failed)


def test_merged_cells_table_emits_table_complex_warning():
    from distill.ir import Table, TableRow, TableCell, Paragraph, TextRun
    from distill.renderer import MarkdownRenderer
    from distill.parsers.base import ParseOptions

    collector = WarningCollector()
    options = ParseOptions(collector=collector)

    table = Table(
        rows=[
            TableRow(cells=[
                TableCell(content=[Paragraph(runs=[TextRun(text='A')])]),
                TableCell(content=[Paragraph(runs=[TextRun(text='B')])]),
            ]),
        ],
        merged_cells=True,
    )

    renderer = MarkdownRenderer()
    renderer._render_table(table, options=options)

    assert collector.has(WarningType.table_complex)


def test_convert_result_has_structured_warnings_field():
    from distill import convert

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(_make_docx_bytes())
        tmp = Path(f.name)

    try:
        result = convert(tmp)
    finally:
        tmp.unlink(missing_ok=True)

    assert hasattr(result, "structured_warnings")
    assert isinstance(result.structured_warnings, list)


def test_convert_clean_document_returns_empty_structured_warnings():
    from distill import convert

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(_make_docx_bytes())
        tmp = Path(f.name)

    try:
        result = convert(tmp)
    finally:
        tmp.unlink(missing_ok=True)

    assert result.structured_warnings == []
